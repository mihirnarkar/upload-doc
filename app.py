from flask import Flask, request, render_template, session, redirect, url_for
from werkzeug.utils import secure_filename
import os
from docx import Document
import fitz  # PyMuPDF library for PDF processing
from docx.opc.exceptions import PackageNotFoundError

app = Flask(__name__)
app.secret_key = 'your_secret_key'

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'doc', 'docx', 'pdf'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_docx(filepath):
    try:
        doc = Document(filepath)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    except PackageNotFoundError:
        return "Invalid Word document"

def read_pdf(filepath):
    with fitz.open(filepath) as pdf_document:
        return '\n'.join([page.get_text() for page in pdf_document])

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            return render_template('index.html', message='No file part')

        files = request.files.getlist('file')
        uploaded_files = session.get('uploaded_files', [])

        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                uploaded_files.append(filename)

        session['uploaded_files'] = uploaded_files

        return render_template('index.html', message='Files uploaded successfully!', uploaded_files=uploaded_files)

    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        return 'File uploaded successfully!'
    else:
        return 'Invalid file type'

@app.route('/submit', methods=['POST'])
def submit():
    combined_text = ""

    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        if filename.endswith(('.doc', '.docx')):
            doc = Document(filepath)
            for paragraph in doc.paragraphs:
                combined_text += paragraph.text + '\n'

        elif filename.endswith('.pdf'):
            with fitz.open(filepath) as pdf_document:
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    combined_text += page.get_text() + '\n'

    with open(os.path.join(app.config['UPLOAD_FOLDER'], 'combined_text.txt'), 'w', encoding='utf-8') as output_file:
        output_file.write(combined_text)

    return render_template('index.html', message='Text file created successfully!')

if __name__ == '__main__':
    app.run(debug=True)
